from __future__ import annotations

import io
import os
import re
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image
from sqlalchemy import select
from sqlalchemy.orm import Session

from .config import Settings
from .models import (
    Benefit,
    BrandingProfile,
    Capability,
    CapabilityMapping,
    Engagement,
    EvidenceItem,
    FileObject,
    Finding,
    PromptDefinition,
    Prospect,
    Report,
    ReportSection,
    Response,
    SectionContentVersion,
    Site,
    User,
)
from .storage import ObjectStorage, StorageConfigurationError


DEFAULT_PROPRIETARY_FOOTER = (
    "This document is the property of and proprietary to Cloud Inventory and contains trade secret and "
    "confidential information, and is solely for the Customer's internal use. Without the express written "
    "consent of Cloud Inventory, this document shall not be used, reproduced, copied, disclosed, or transmitted, "
    "in whole or in part. Copyright Cloud Inventory. All rights reserved."
)


def _hex_color(value: str) -> RGBColor:
    value = value.lstrip("#")
    return RGBColor.from_string(value.upper())


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.lstrip("#"))
    tc_pr.append(shd)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _add_field(paragraph, instruction: str, placeholder: str = "", *, dirty: bool = False) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    if dirty:
        begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def _request_field_updates(doc: Document) -> None:
    settings = doc.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Page ")
    _add_field(paragraph, "PAGE", "1")


def _add_watermark_to_header(header, text: str) -> None:
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape_id = "PowerPlusWaterMarkObject"
    xml = f'''
    <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office">
      <w:rPr><w:noProof/></w:rPr>
      <w:pict>
        <v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" adj="10800" path="m@7,l@8,m@5,21600l@6,21600e">
          <v:formulas><v:f eqn="sum #0 0 10800"/><v:f eqn="prod #0 2 1"/><v:f eqn="sum 21600 0 @1"/><v:f eqn="sum 0 0 @2"/><v:f eqn="sum 21600 0 @3"/><v:f eqn="if @0 @3 0"/><v:f eqn="if @0 21600 @1"/><v:f eqn="if @0 0 @2"/><v:f eqn="if @0 @4 21600"/><v:f eqn="mid @5 @6"/><v:f eqn="mid @8 @5"/><v:f eqn="mid @7 @8"/><v:f eqn="mid @6 @7"/><v:f eqn="sum @6 0 @5"/></v:formulas>
          <v:path textpathok="t" o:connecttype="custom" o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800" o:connectangles="270,180,90,0"/>
          <v:textpath on="t" fitshape="t"/>
          <o:lock v:ext="edit" text="t" shapetype="t"/>
        </v:shapetype>
        <v:shape id="{shape_id}" o:spid="_x0000_s2049" type="#_x0000_t136"
          style="position:absolute;margin-left:0;margin-top:0;width:500pt;height:100pt;rotation:315;z-index:-251654144;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin"
          fillcolor="silver" stroked="f">
          <v:fill opacity="0.18"/>
          <v:textpath style="font-family:'Arial';font-size:1pt" string="{text}"/>
        </v:shape>
      </w:pict>
    </w:r>'''
    paragraph._p.append(parse_xml(xml))


def _configure_styles(doc: Document, brand: BrandingProfile) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = brand.body_font
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    for idx, size in [(1, 18), (2, 14), (3, 12)]:
        style = styles[f"Heading {idx}"]
        style.font.name = brand.heading_font
        style.font.color.rgb = _hex_color(brand.primary_color if idx == 1 else brand.secondary_color)
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
    # Keep list markers visually subordinate to body text and clearly indented.
    for list_style_name in ("List Bullet", "List Number"):
        list_style = styles[list_style_name]
        list_style.font.name = brand.body_font
        list_style.font.size = Pt(10)
        list_style.paragraph_format.left_indent = Inches(0.35)
        list_style.paragraph_format.first_line_indent = Inches(-0.18)
        list_style.paragraph_format.space_after = Pt(3)

    toc_heading = styles["TOC Heading"]
    toc_heading.font.name = brand.heading_font
    toc_heading.font.size = Pt(18)
    toc_heading.font.bold = True
    toc_heading.font.color.rgb = _hex_color(brand.primary_color)
    toc_heading.paragraph_format.space_after = Pt(12)

    if "Evidence Caption" not in [s.name for s in styles]:
        cap = styles.add_style("Evidence Caption", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = brand.body_font
        cap.font.size = Pt(8)
        cap.font.italic = True
        cap.font.color.rgb = _hex_color(brand.accent_color)


def _configure_section_layout(section) -> None:
    section.top_margin = Inches(0.7)
    # Reserve enough body clearance for the legal footer on content pages.
    section.bottom_margin = Inches(0.95)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.footer_distance = Inches(0.22)


def _clear_footer(footer) -> None:
    for table in list(footer.tables):
        footer._element.remove(table._element)
    for paragraph in footer.paragraphs:
        paragraph.clear()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)


def _add_content_footer(section, brand: BrandingProfile, footer_logo_path: Path | None) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    _clear_footer(footer)

    table = footer.add_table(rows=1, cols=3, width=Inches(7.0))
    table.autofit = False
    widths = (Inches(0.82), Inches(5.48), Inches(0.70))
    for column, width in zip(table.columns, widths):
        column.width = width
    for cell, width in zip(table.rows[0].cells, widths):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Set explicit OOXML grid widths so LibreOffice preserves the asymmetric
    # footer layout while refreshing the TOC field.
    grid_cols = table._tbl.tblGrid.gridCol_lst
    for grid_col, width in zip(grid_cols, widths):
        grid_col.set(qn("w:w"), str(int(width.twips)))

    logo_cell, text_cell, page_cell = table.rows[0].cells
    logo_p = logo_cell.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_p.paragraph_format.space_after = Pt(0)
    if footer_logo_path and footer_logo_path.exists():
        logo_p.add_run().add_picture(str(footer_logo_path), width=Inches(0.68))

    legal_p = text_cell.paragraphs[0]
    legal_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    legal_p.paragraph_format.space_after = Pt(0)
    legal_run = legal_p.add_run(brand.footer_text or DEFAULT_PROPRIETARY_FOOTER)
    legal_run.font.name = brand.body_font
    legal_run.font.size = Pt(5.5)
    legal_run.font.color.rgb = RGBColor(70, 78, 86)

    page_p = page_cell.paragraphs[0]
    page_p.paragraph_format.space_after = Pt(0)
    _add_page_number(page_p)
    for run in page_p.runs:
        run.font.name = brand.body_font
        run.font.size = Pt(7)


def _apply_headers_footers(
    doc: Document,
    brand: BrandingProfile,
    is_final: bool,
    footer_logo_path: Path | None,
) -> None:
    if not doc.sections:
        return

    # The cover occupies its own section so page 1 is structurally footer-free.
    # This survives LibreOffice's TOC refresh, which can otherwise duplicate a
    # default footer into a first-page footer even when w:titlePg is present.
    cover_section = doc.sections[0]
    _configure_section_layout(cover_section)
    cover_section.footer.is_linked_to_previous = False
    _clear_footer(cover_section.footer)
    cover_section.first_page_footer.is_linked_to_previous = False
    _clear_footer(cover_section.first_page_footer)

    if not is_final:
        cover_section.header.is_linked_to_previous = False
        _add_watermark_to_header(cover_section.header, brand.draft_watermark)

    for section in doc.sections[1:]:
        _configure_section_layout(section)
        _add_content_footer(section, brand, footer_logo_path)
        if not is_final:
            section.header.is_linked_to_previous = False
            _add_watermark_to_header(section.header, brand.draft_watermark)


def _add_logo(doc: Document, logo_path: Path | None, width: float = 2.0) -> None:
    if logo_path and logo_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo_path), width=Inches(width))


def _add_cover(doc: Document, report: Report, prospect: Prospect, site: Site | None, engagement: Engagement, brand: BrandingProfile, logo_path: Path | None, is_final: bool) -> None:
    _add_logo(doc, logo_path, 2.3)
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(report.title)
    run.bold = True
    run.font.name = brand.heading_font
    run.font.size = Pt(28)
    run.font.color.rgb = _hex_color(brand.primary_color)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Site Discovery Report").bold = True
    subtitle.runs[0].font.size = Pt(18)
    doc.add_paragraph()
    info = doc.add_table(rows=4, cols=2)
    info.style = "Table Grid"
    values = [
        ("Prospect", prospect.name),
        ("Survey location", site.name + (f" - {site.address}" if site and site.address else "") if site else "Not specified"),
        ("Survey date", engagement.survey_date.isoformat() if engagement.survey_date else "Not specified"),
        ("Document status", "FINAL" if is_final else "DRAFT - CONFIDENTIAL"),
    ]
    for row, (label, value) in zip(info.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        _set_cell_shading(row.cells[0], brand.primary_color)
        for r in row.cells[0].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.bold = True
    doc.add_paragraph()
    note = doc.add_paragraph(brand.confidentiality_text)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in note.runs:
        run.font.size = Pt(8)
        run.font.italic = True



def _add_toc(doc: Document) -> None:
    # Word Automatic Table 2 equivalent: built-in TOC Heading plus a real TOC
    # field using Heading 1-3, hyperlinks and page-number references.
    doc.add_paragraph("Table of Contents", style="TOC Heading")
    field_paragraph = doc.add_paragraph()
    _add_field(
        field_paragraph,
        ' TOC \\o "1-3" \\h \\z \\u ',
        "Table of Contents will update automatically.",
        dirty=True,
    )
    doc.add_page_break()


def _add_revision_history(doc: Document, report: Report, owner: User | None, brand: BrandingProfile) -> None:
    doc.add_heading("Document Revision History", level=1)
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    headers = ["Changed By", "Date", "Version", "Notes"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
        _set_cell_shading(table.rows[0].cells[idx], brand.primary_color)
        for run in table.rows[0].cells[idx].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    _set_repeat_table_header(table.rows[0])
    values = [owner.display_name or owner.username if owner else "Report owner", datetime.now(timezone.utc).strftime("%Y-%m-%d"), str(report.revision), "Generated from the Cloud Inventory Site Discovery Platform"]
    for idx, value in enumerate(values):
        table.rows[1].cells[idx].text = value
    doc.add_page_break()


def _add_text(doc: Document, text: str) -> None:
    for block in [x.strip() for x in text.split("\n") if x.strip()]:
        if block.startswith(("- ", "• ", "* ")):
            doc.add_paragraph(block[2:].strip(), style="List Bullet")
            continue
        numbered = re.match(r"^\d+[\.)]\s+(.+)$", block)
        if numbered:
            doc.add_paragraph(numbered.group(1).strip(), style="List Number")
            continue
        doc.add_paragraph(block)


def _image_bytes_to_file(data: bytes, suffix: str = ".jpg") -> Path:
    fd, path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    p = Path(path)
    p.write_bytes(data)
    return p


def _add_evidence(doc: Document, db: Session, storage: ObjectStorage | None, evidence: Iterable[EvidenceItem]) -> None:
    for item in evidence:
        file_obj = db.scalar(select(FileObject).where(FileObject.evidence_id == item.id, FileObject.variant.in_(["WEB", "ORIGINAL"])).order_by(FileObject.variant.desc()))
        if not file_obj:
            continue
        if storage is None:
            doc.add_paragraph(f"Evidence unavailable during generation: {file_obj.file_name} (persistent object storage is not configured)", style="Evidence Caption")
            continue
        try:
            data = storage.get_bytes(file_obj.storage_key)
            if file_obj.mime_type.startswith("image/"):
                suffix = Path(file_obj.file_name).suffix or ".jpg"
                temp = _image_bytes_to_file(data, suffix)
                try:
                    with Image.open(temp) as image:
                        width, height = image.size
                    max_width = 6.7
                    width_inches = min(max_width, max(2.5, max_width if width >= height else 4.2))
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(str(temp), width=Inches(width_inches))
                    cap = doc.add_paragraph(item.caption or file_obj.file_name, style="Evidence Caption")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                finally:
                    temp.unlink(missing_ok=True)
            else:
                doc.add_paragraph(f"Supporting attachment: {item.caption or file_obj.file_name}", style="List Bullet")
        except Exception as exc:
            doc.add_paragraph(f"Evidence unavailable during generation: {file_obj.file_name} ({exc})", style="Evidence Caption")


def generate_docx(db: Session, report_id: str, settings: Settings, *, publication_type: str, is_final: bool) -> bytes:
    report = db.get(Report, report_id)
    if not report:
        raise ValueError("Report not found")
    prospect = db.get(Prospect, report.prospect_id)
    engagement = db.get(Engagement, report.engagement_id)
    site = db.get(Site, report.site_id) if report.site_id else None
    brand = db.get(BrandingProfile, report.branding_profile_id) if report.branding_profile_id else db.scalar(select(BrandingProfile).where(BrandingProfile.is_default.is_(True), BrandingProfile.active.is_(True)))
    if not prospect or not engagement or not brand:
        raise ValueError("Report dependencies are incomplete")
    owner = db.get(User, report.owner_id)
    storage: ObjectStorage | None = None
    try:
        storage = ObjectStorage(settings)
    except StorageConfigurationError:
        # Draft generation must remain available even when persistent R2 storage
        # has not yet been configured. Stored evidence/custom branding is simply
        # omitted with an explanatory note instead of aborting the document.
        storage = None
    cloud_inventory_logo_path = Path(__file__).parent / "static" / "cloud-inventory-logo-for-light-background-v0.4.1.png"
    logo_path = cloud_inventory_logo_path
    custom_logo_path: Path | None = None
    if brand.logo_storage_key and storage is not None:
        try:
            fd, custom_name = tempfile.mkstemp(prefix="ci-discovery-logo-", suffix=".png")
            os.close(fd)
            custom_logo_path = Path(custom_name)
            custom_logo_path.write_bytes(storage.get_bytes(brand.logo_storage_key))
            logo_path = custom_logo_path
        except Exception:
            if custom_logo_path:
                custom_logo_path.unlink(missing_ok=True)
            custom_logo_path = None

    doc = Document()
    _configure_styles(doc, brand)
    _request_field_updates(doc)
    _add_cover(doc, report, prospect, site, engagement, brand, logo_path, is_final)
    # Start page 2 in a separate section so the footer is guaranteed to appear
    # on every page after the cover, but never on page 1.
    doc.add_section(WD_SECTION.NEW_PAGE)
    _apply_headers_footers(doc, brand, is_final, cloud_inventory_logo_path)
    _add_revision_history(doc, report, owner, brand)

    if publication_type == "FOLLOW_UP_QUESTIONNAIRE":
        doc.add_heading("Customer Follow-up Questionnaire", level=1)
    elif publication_type == "DEMO_BRIEF":
        doc.add_heading("Solution Demonstration Brief", level=1)

    all_sections = list(db.scalars(select(ReportSection).where(ReportSection.report_id == report.id, ReportSection.state != "REMOVED").order_by(ReportSection.display_order)).all())

    def has_publishable_content(section: ReportSection) -> bool:
        if section.narrative.strip():
            return True
        checks = [
            select(Response.id).where(Response.section_id == section.id).limit(1),
            select(Finding.id).where(Finding.section_id == section.id, Finding.status != "REJECTED").limit(1),
            select(EvidenceItem.id).where(EvidenceItem.section_id == section.id, EvidenceItem.status.in_(["READY", "AVAILABLE"])).limit(1),
            select(SectionContentVersion.id).where(
                SectionContentVersion.section_id == section.id,
                SectionContentVersion.content_type == "CLOUD_INVENTORY_APPROACH",
                SectionContentVersion.is_current.is_(True),
            ).limit(1),
        ]
        return any(db.scalar(stmt) is not None for stmt in checks)

    if publication_type == "FOLLOW_UP_QUESTIONNAIRE":
        sections = all_sections
    elif publication_type == "DEMO_BRIEF":
        sections = [s for s in all_sections if has_publishable_content(s) or s.stable_key in {"executive-summary", "vision-pain-points", "solution-viability", "expected-benefits", "next-steps"}]
    else:
        # Draft and final discovery documents include only sections that contain
        # reportable content. Empty optional sections never create blank output.
        sections = [s for s in all_sections if has_publishable_content(s)]

    _add_toc(doc)
    for section in sections:
        doc.add_heading(section.title, level=1)
        if publication_type == "FOLLOW_UP_QUESTIONNAIRE":
            prompts = list(db.scalars(select(PromptDefinition).where(PromptDefinition.active.is_(True), (PromptDefinition.process_module == section.process_module) if section.process_module else PromptDefinition.process_module.is_(None)).order_by(PromptDefinition.display_order)).all())
            answered = {r.prompt_id for r in db.scalars(select(Response).where(Response.section_id == section.id)).all()}
            unanswered = [p for p in prompts if p.id not in answered]
            if unanswered:
                for prompt in unanswered:
                    doc.add_paragraph(prompt.question, style="List Number")
                    doc.add_paragraph("Response: ______________________________________________________________")
            else:
                doc.add_paragraph("No outstanding structured questions were identified for this section.")
            continue

        if section.narrative.strip():
            _add_text(doc, section.narrative)
        responses = db.execute(select(Response, PromptDefinition).join(PromptDefinition, Response.prompt_id == PromptDefinition.id).where(Response.section_id == section.id, PromptDefinition.active.is_(True)).order_by(PromptDefinition.display_order)).all()
        if responses:
            doc.add_heading("Discovery Responses", level=2)
            for response, prompt in responses:
                p = doc.add_paragraph()
                p.add_run(prompt.question).bold = True
                _add_text(doc, response.narrative or (str(response.payload) if response.payload else "No narrative response recorded."))

        findings = list(db.scalars(select(Finding).where(Finding.section_id == section.id, Finding.status != "REJECTED").order_by(Finding.created_at)).all())
        if findings:
            doc.add_heading("Current-State Findings", level=2)
            for finding in findings:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{finding.finding_type.replace('_', ' ').title()}: ").bold = True
                p.add_run(finding.statement)
                if finding.impact:
                    doc.add_paragraph(f"Impact: {finding.impact}")

        solution = db.scalar(
            select(SectionContentVersion)
            .where(
                SectionContentVersion.report_id == report.id,
                SectionContentVersion.section_id == section.id,
                SectionContentVersion.content_type == "CLOUD_INVENTORY_APPROACH",
                SectionContentVersion.is_current.is_(True),
            )
            .order_by(SectionContentVersion.version.desc())
        )
        if solution and solution.text.strip():
            doc.add_heading("Cloud Inventory Approach", level=2)
            _add_text(doc, solution.text)

        mappings = db.execute(
            select(CapabilityMapping, Capability)
            .join(Capability, CapabilityMapping.capability_id == Capability.id)
            .where(
                CapabilityMapping.report_id == report.id,
                CapabilityMapping.section_id == section.id,
                CapabilityMapping.approval_state == "APPROVED",
            )
            .order_by(Capability.name)
        ).all()
        if mappings:
            doc.add_heading("Mapped Cloud Inventory Functionality", level=3 if solution and solution.text.strip() else 2)
            for mapping, capability in mappings:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(capability.name + ": ").bold = True
                p.add_run(mapping.rationale)
                if mapping.source_label and mapping.source_statement:
                    source_p = doc.add_paragraph()
                    source_p.add_run(f"Mapped from {mapping.source_label}: ").bold = True
                    source_p.add_run(mapping.source_statement)
                if mapping.prerequisites or capability.typical_prerequisites:
                    doc.add_paragraph("Prerequisites: " + (mapping.prerequisites or capability.typical_prerequisites or ""))
                if capability.limitations:
                    doc.add_paragraph("Limitations: " + capability.limitations)

        benefits = list(db.scalars(select(Benefit).join(Finding, Benefit.finding_id == Finding.id, isouter=True).where(Benefit.report_id == report.id, Benefit.approval_state == "APPROVED", (Finding.section_id == section.id) | (Benefit.finding_id.is_(None)))).all())
        if benefits:
            doc.add_heading("Benefits", level=2)
            for benefit in benefits:
                doc.add_paragraph(benefit.statement, style="List Bullet")
                if benefit.measure_type == "QUANTITATIVE" and (benefit.formula or benefit.assumptions):
                    doc.add_paragraph(f"Measurement basis: {benefit.formula or 'To be established'}. Assumptions: {benefit.assumptions or 'None recorded.'}")

        section_evidence = list(db.scalars(select(EvidenceItem).where(EvidenceItem.section_id == section.id, EvidenceItem.status.in_(["READY", "AVAILABLE"]), EvidenceItem.placement == "INLINE").order_by(EvidenceItem.created_at)).all())
        if section_evidence:
            doc.add_heading("Site Photographs and Evidence", level=2)
            _add_evidence(doc, db, storage, section_evidence)

    appendix = list(db.scalars(select(EvidenceItem).where(EvidenceItem.report_id == report.id, EvidenceItem.status.in_(["READY", "AVAILABLE"]), EvidenceItem.placement == "APPENDIX").order_by(EvidenceItem.created_at)).all())
    if appendix and publication_type != "FOLLOW_UP_QUESTIONNAIRE":
        doc.add_page_break()
        doc.add_heading("Appendix - Supporting Evidence", level=1)
        _add_evidence(doc, db, storage, appendix)

    output = io.BytesIO()
    doc.save(output)
    if custom_logo_path:
        custom_logo_path.unlink(missing_ok=True)
    return output.getvalue()


def refresh_docx_fields(
    docx_bytes: bytes,
    settings: Settings,
    *,
    emit_pdf: bool = False,
) -> tuple[bytes, bytes | None]:
    """Refresh TOC/page-reference fields through LibreOffice UNO when available.

    Render runs use a Linux system Python with python3-uno. Windows/local
    validation falls back to the original DOCX; the embedded TOC field remains
    valid and is marked for automatic refresh when opened in Word.
    """
    settings.document_work_dir.mkdir(parents=True, exist_ok=True)
    helper = Path(__file__).resolve().parents[1] / "scripts" / "refresh_docx_fields.py"
    system_python = Path("/usr/bin/python3")
    if os.name == "nt" or not helper.exists() or not system_python.exists():
        return docx_bytes, None

    with tempfile.TemporaryDirectory(dir=settings.document_work_dir) as tmp:
        tmp_path = Path(tmp)
        docx_path = tmp_path / "report.docx"
        pdf_path = tmp_path / "report.pdf"
        docx_path.write_bytes(docx_bytes)
        cmd = [
            str(system_python),
            str(helper),
            str(docx_path),
            "--libreoffice",
            settings.libreoffice_path,
        ]
        if emit_pdf:
            cmd.extend(["--pdf", str(pdf_path)])
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=210)
        if proc.returncode != 0 or not docx_path.exists():
            return docx_bytes, None
        refreshed = docx_path.read_bytes()
        pdf_bytes = pdf_path.read_bytes() if emit_pdf and pdf_path.exists() else None
        return refreshed, pdf_bytes


def convert_docx_to_pdf(docx_bytes: bytes, settings: Settings) -> bytes:
    settings.document_work_dir.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(dir=settings.document_work_dir) as tmp:
        tmp_path = Path(tmp)
        docx_path = tmp_path / "report.docx"
        docx_path.write_bytes(docx_bytes)
        profile = tmp_path / "lo-profile"
        profile.mkdir()
        cmd = [
            settings.libreoffice_path,
            f"-env:UserInstallation={profile.resolve().as_uri()}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(tmp_path),
            str(docx_path),
        ]
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
        pdf_path = tmp_path / "report.pdf"
        if proc.returncode != 0 or not pdf_path.exists():
            raise RuntimeError(f"LibreOffice conversion failed: {proc.stderr or proc.stdout}")
        return pdf_path.read_bytes()
