from __future__ import annotations

import csv
import io
import zipfile
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

MAX_EXTRACTED_CHARACTERS = 100_000


class ExtractionResult:
    def __init__(self, state: str, text: str | None = None, reason: str | None = None):
        self.state = state
        self.text = text
        self.reason = reason


def _truncate(value: str) -> str:
    value = value.replace("\x00", "").strip()
    if len(value) <= MAX_EXTRACTED_CHARACTERS:
        return value
    return value[:MAX_EXTRACTED_CHARACTERS] + "\n\n[Extraction truncated]"


def _extract_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                parts.append(" | ".join(values))
    return "\n".join(parts)


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[Page {index}]\n{text}")
    return "\n\n".join(pages)


def _extract_xlsx(data: bytes) -> str:
    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    output: list[str] = []
    for sheet in workbook.worksheets:
        output.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(value.strip() for value in values):
                output.append(" | ".join(values))
    workbook.close()
    return "\n".join(output)


def _extract_csv(data: bytes) -> str:
    text = data.decode("utf-8-sig", errors="replace")
    rows = csv.reader(io.StringIO(text))
    return "\n".join(" | ".join(row) for row in rows)


def extract_text(data: bytes, filename: str, mime_type: str | None) -> ExtractionResult:
    """Extract searchable text from safe, supported office/text formats.

    Images intentionally return NOT_APPLICABLE: OCR is not performed automatically because it
    is expensive and error-prone. AI image review may be added later behind the same approval gate.
    """
    suffix = Path(filename).suffix.lower()
    mime = (mime_type or "").lower()
    try:
        if mime.startswith("text/") or suffix in {".txt", ".md", ".log", ".json", ".xml"}:
            text = data.decode("utf-8-sig", errors="replace")
        elif suffix == ".csv" or mime == "text/csv":
            text = _extract_csv(data)
        elif suffix == ".docx" or mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = _extract_docx(data)
        elif suffix == ".pdf" or mime == "application/pdf":
            text = _extract_pdf(data)
        elif suffix in {".xlsx", ".xlsm"} or mime == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            text = _extract_xlsx(data)
        elif suffix == ".zip" or mime in {"application/zip", "application/x-zip-compressed"}:
            # Do not unpack arbitrary ZIP payloads. Record only a safe file inventory.
            with zipfile.ZipFile(io.BytesIO(data)) as archive:
                text = "[ZIP file inventory]\n" + "\n".join(archive.namelist()[:500])
        else:
            return ExtractionResult("NOT_APPLICABLE", reason="Unsupported extraction format")
        cleaned = _truncate(text)
        return ExtractionResult("COMPLETED" if cleaned else "EMPTY", cleaned or None)
    except Exception as exc:  # A failed extraction must never fail the underlying upload.
        return ExtractionResult("FAILED", reason=str(exc)[:1000])
