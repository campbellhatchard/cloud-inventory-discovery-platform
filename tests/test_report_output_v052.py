from __future__ import annotations

import io
import re
import zipfile

from docx import Document

from app.config import get_settings
from app.database import SessionLocal
from app.documents import generate_docx
from app.models import Publication


LEGAL_FOOTER = (
    "This document is the property of and proprietary to Cloud Inventory and contains trade secret and "
    "confidential information, and is solely for the Customer's internal use. Without the express written "
    "consent of Cloud Inventory, this document shall not be used, reproduced, copied, disclosed, or transmitted, "
    "in whole or in part. Copyright Cloud Inventory. All rights reserved."
)


def headers(me: dict) -> dict[str, str]:
    return {"X-CSRF-Token": me["csrf_token"]}


def create_report(client, me: dict) -> str:
    h = headers(me)
    onboard = client.post(
        "/api/prospects/onboard",
        json={
            "prospect": {"name": "Report Output Test"},
            "site": {"name": "Main Site", "timezone": "America/Denver"},
            "engagement": {"name": "Discovery"},
        },
        headers=h,
    )
    assert onboard.status_code == 200, onboard.text
    templates = client.get("/api/report-templates").json()
    report = client.post(
        f"/api/prospects/{onboard.json()['id']}/reports",
        json={
            "title": "Report Output Formatting Test",
            "engagement_id": onboard.json()["engagement_id"],
            "report_template_id": templates[0]["id"],
            "report_kind": "CONSOLIDATED",
        },
        headers=h,
    )
    assert report.status_code == 200, report.text
    return report.json()["id"]


def test_docx_uses_real_toc_indented_lists_and_legal_footer(admin_session) -> None:
    client, me = admin_session
    report_id = create_report(client, me)
    payload = client.get(f"/api/reports/{report_id}").json()
    section = next(item for item in payload["sections"] if item["stable_key"] == "picking")
    update = client.patch(
        f"/api/reports/{report_id}/sections/{section['id']}",
        json={
            "narrative": "Picking observations\n- Paper pick tickets are released in waves.\n1. Pickers collect the paper ticket.\n2) Completed tickets return to the desk.",
            "expected_version": section["version"],
        },
        headers=headers(me),
    )
    assert update.status_code == 200, update.text

    with SessionLocal() as db:
        data = generate_docx(
            db,
            report_id,
            get_settings(),
            publication_type="FULL_DISCOVERY",
            is_final=False,
        )

    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
        footer_xml = "\n".join(
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        )

        assert "Table of Contents" in document_xml
        assert 'TOC \\o "1-3" \\h \\z \\u' in document_xml.replace("&quot;", '"')
        assert "updateFields" in settings_xml
        assert "ListBullet" in document_xml
        assert "ListNumber" in document_xml
        assert re.search(r'w:styleId="ListBullet".*?<w:ind[^>]+w:left="504"[^>]+w:hanging="259"', styles_xml, re.S)
        assert re.search(r'w:styleId="ListNumber".*?<w:ind[^>]+w:left="504"[^>]+w:hanging="259"', styles_xml, re.S)
        assert document_xml.count("<w:sectPr") >= 2
        assert LEGAL_FOOTER in footer_xml
        assert "Cloud Inventory | Confidential" not in footer_xml
        assert "PAGE" in footer_xml
        assert "<w:drawing>" in footer_xml

    rendered_doc = Document(io.BytesIO(data))
    assert len(rendered_doc.sections) >= 2
    assert not rendered_doc.sections[0].footer.tables
    assert not "".join(p.text for p in rendered_doc.sections[0].footer.paragraphs).strip()
    assert rendered_doc.sections[1].footer.tables
    assert LEGAL_FOOTER in "".join(
        cell.text
        for table in rendered_doc.sections[1].footer.tables
        for row in table.rows
        for cell in row.cells
    )


def test_failed_publication_can_be_dismissed_without_deleting_history(admin_session) -> None:
    client, me = admin_session
    report_id = create_report(client, me)

    with SessionLocal() as db:
        failed = Publication(
            report_id=report_id,
            report_revision=7,
            publication_type="FULL_DISCOVERY",
            is_final=False,
            status="FAILED",
            error="Historical test failure",
            requested_by=me["id"],
        )
        db.add(failed)
        db.commit()
        publication_id = failed.id

    before = client.get(f"/api/reports/{report_id}")
    assert before.status_code == 200
    item = next(p for p in before.json()["publications"] if p["id"] == publication_id)
    assert item["report_revision"] == 7
    assert item["created_at"]

    dismissed = client.post(
        f"/api/reports/{report_id}/publications/{publication_id}/dismiss",
        headers=headers(me),
    )
    assert dismissed.status_code == 200, dismissed.text
    assert dismissed.json()["dismissed_at"]

    after = client.get(f"/api/reports/{report_id}")
    assert publication_id not in {p["id"] for p in after.json()["publications"]}

    with SessionLocal() as db:
        stored = db.get(Publication, publication_id)
        assert stored is not None
        assert stored.dismissed_at is not None
        assert stored.dismissed_by == me["id"]


def test_v052_frontend_and_migration_contract() -> None:
    root = __import__("pathlib").Path(__file__).resolve().parents[1]
    app_js = (root / "app" / "static" / "app.js").read_text(encoding="utf-8")
    migration = (
        root
        / "alembic"
        / "versions"
        / "b72e1f8c5d21_report_output_publication_history.py"
    ).read_text(encoding="utf-8")
    dockerfile = (root / "Dockerfile").read_text(encoding="utf-8")

    assert "PREVIOUS FAILED ATTEMPT" in app_js
    assert "Dismiss failed attempt" in app_js
    assert "fmtDateTime" in app_js
    assert "/dismiss" in app_js
    assert "dismissed_at" in migration
    assert "dismissed_by" in migration
    assert "Cloud Inventory | Confidential" in migration
    assert "python3-uno" in dockerfile
